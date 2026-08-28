#!/usr/bin/env python3
"""Build the plain-English David Leads broker field guide."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("David_Leads_Broker_Guide.docx")

NAVY = "0A2540"
TEAL = "168F89"
GOLD = "B78327"
INK = "1F2933"
MUTED = "5A6B7C"
LINE = "D7E0E8"
PALE_TEAL = "E8F6F4"
PALE_BLUE = "E8EEF5"
PALE_GOLD = "FFF7E6"
WHITE = "FFFFFF"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM = 80
CELL_START_END = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_TOP_BOTTOM),
        ("bottom", CELL_TOP_BOTTOM),
        ("start", CELL_START_END),
        ("end", CELL_START_END),
    ):
        element = tc_mar.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def mark_header_row(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    row_properties.append(marker)


def prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run(run, *, size: float | None = None, bold: bool | None = None,
            italic: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), TEAL)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    if bold:
        run_properties.append(OxmlElement("w:b"))
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22")
    run_properties.append(size)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_numbering(document: Document, *, bullet: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "\u2022" if bullet else "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    paragraph_properties = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    paragraph_properties.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    paragraph_properties.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    paragraph_properties.append(spacing)
    level.append(paragraph_properties)
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    run_properties.append(fonts)
    level.append(run_properties)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    number.append(abstract_ref)
    numbering.append(number)
    return num_id


def add_list_item(
    document: Document,
    text: str,
    *,
    marker: str = "•",
    lead: str | None = None,
):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(f"{marker} ")
    set_run(run, bold=marker != "•", color=NAVY)
    if lead:
        run = paragraph.add_run(lead)
        set_run(run, bold=True, color=NAVY)
    run = paragraph.add_run(text)
    set_run(run, color=INK)
    return paragraph


def add_body(document: Document, text: str, *, bold_lead: str | None = None,
             italic: bool = False, color: str = INK):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead:
        run = paragraph.add_run(bold_lead)
        set_run(run, bold=True, color=NAVY)
    run = paragraph.add_run(text)
    set_run(run, italic=italic, color=color)
    return paragraph


def add_callout(document: Document, label: str, text: str, *, fill: str = PALE_TEAL):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    properties.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), TEAL)
    borders.append(left)
    properties.append(borders)
    indentation = OxmlElement("w:ind")
    indentation.set(qn("w:left"), "180")
    indentation.set(qn("w:right"), "120")
    properties.append(indentation)
    run = paragraph.add_run(label + " ")
    set_run(run, bold=True, color=NAVY)
    run = paragraph.add_run(text)
    set_run(run, color=INK)
    return paragraph


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run(run, size=9, color=MUTED)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(13.5)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_after = Pt(18)

    tokens = {
        "Heading 1": (16, NAVY, 18, 10),
        "Heading 2": (13, TEAL, 14, 7),
        "Heading 3": (12, NAVY, 10, 5),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_page(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    run = header_paragraph.add_run("DAVID LEADS  |  BROKER FIELD GUIDE")
    set_run(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer_paragraph.add_run("Operating guide  |  verify the current release at demo time  |  ")
    set_run(run, size=8.5, color=MUTED)
    add_page_number(footer_paragraph)


def set_table_text(cell, text: str, *, bold: bool = False, color: str = INK,
                   size: float = 10.5) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run(run, bold=bold, color=color, size=size)


def build() -> Document:
    document = Document()
    configure_styles(document)
    configure_page(document)
    document.core_properties.title = "David Leads Broker Field Guide"
    document.core_properties.subject = "Plain-English operating guide for David Abraham"
    document.core_properties.author = "SZL Holdings"
    document.core_properties.keywords = "David Leads, broker, public data, Eastern United States"
    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("BROKER ENABLEMENT PACK")
    set_run(run, size=9.5, bold=True, color=GOLD)
    title = document.add_paragraph(style="Title")
    title.add_run("David Leads")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("A plain-English field guide for Evidence-Backed Broker Research")
    prepared = document.add_paragraph()
    prepared.paragraph_format.space_after = Pt(18)
    run = prepared.add_run("Prepared for David Abraham  |  Live state must be verified at demo time")
    set_run(run, size=10.5, bold=True, color=MUTED)

    add_callout(
        document,
        "Start here:",
        "Viewing the live research workspace does not require a username, password, or access key. Open the web address below in Chrome, Edge, Safari, or a phone browser.",
    )
    link_paragraph = document.add_paragraph()
    link_paragraph.paragraph_format.space_after = Pt(16)
    add_hyperlink(
        link_paragraph,
        "Open David Leads",
        "https://szlholdings-david-leads.hf.space/",
        bold=True,
    )

    add_heading(document, "What David Leads does", 1)
    add_body(
        document,
        "It gathers current organization-level records from official public sources, shows the business event that changed, explains why the timing may deserve research, and links back to the source record. It is a research desk, not a purchased contact list.",
    )
    add_list_item(document, "Choose all 27 Eastern markets, a region, or one state.")
    add_list_item(document, "See live business and facility records with source links and receipts.")
    add_list_item(document, "Focus on the strongest timing signals before doing manual research.")
    add_list_item(document, "Keep every outreach decision behind a human compliance check.")

    add_heading(document, "Live release checklist", 2)
    add_body(
        document,
        "Do not reuse a dated count or source state. Read each value from the running release immediately before the demo.",
        italic=True,
        color=MUTED,
    )
    table = document.add_table(rows=1, cols=2)
    set_table_geometry(table, [2700, 6660])
    mark_header_row(table.rows[0])
    set_cell_shading(table.cell(0, 0), NAVY)
    set_cell_shading(table.cell(0, 1), NAVY)
    set_table_text(table.cell(0, 0), "Check", bold=True, color=WHITE)
    set_table_text(table.cell(0, 1), "Required evidence", bold=True, color=WHITE)
    verified_rows = [
        ("Release identity", "Match /api/build-info to the intended GitHub revision and release attestation."),
        ("Source health", "Read every LIVE, UNAVAILABLE, or NOT_APPLICABLE state from the current Market coverage panel."),
        ("Records and samples", "Treat the count as a current-pull observation; the active path must return no sample records."),
        ("Readiness", "Verify /healthz and /readyz; do not infer readiness from an HTTP 200 alone."),
    ]
    for index, (label, value) in enumerate(verified_rows, start=1):
        cells = table.add_row().cells
        if index % 2 == 0:
            set_cell_shading(cells[0], "F5F8FB")
            set_cell_shading(cells[1], "F5F8FB")
        set_table_text(cells[0], label, bold=True, color=NAVY)
        set_table_text(cells[1], value)
    for row in table.rows:
        prevent_row_split(row)
    set_table_geometry(table, [2700, 6660])

    add_heading(document, "Your first five minutes", 1)
    add_list_item(document, "Open the link. The research workspace loads without a login.", marker="1.")
    add_list_item(document, "Choose All East, a region, or one state. One state runs a deeper pull.", marker="2.")
    add_list_item(document, "Read the Broker brief near the top. It points to the strongest current timing pattern.", marker="3.")
    add_list_item(document, "Use the deal-moment buttons to narrow the list: life-plan timing, growth and awards, operational change, or needs research.", marker="4.")
    add_list_item(document, "Open an organization to see the public event, likely fit, evidence, limits, and the next research step.", marker="5.")
    add_list_item(document, "Open the official source before using any claim in a conversation.", marker="6.")

    add_heading(document, "What the screen is telling you", 2)
    add_body(document, "The number of official records in the selected market. The count changes with the territory and source response.", bold_lead="Live organizations. ")
    add_body(document, "How many source lanes answered live. An unavailable source remains visible; it is never replaced by a fake record.", bold_lead="Source coverage. ")
    add_body(document, "Records worth checking. This is not the number of people cleared for contact.", bold_lead="Needs research. ")
    add_body(document, "A date or filing field that may make the organization worth researching now. It is not proof of a sale, renewal, dissatisfaction, or insurance need.", bold_lead="Timing window. ")

    add_callout(
        document,
        "The key distinction:",
        "A visible public record is permission to research the organization. It is not permission to call, text, email, quote, or recommend a product.",
        fill=PALE_GOLD,
    )

    document.add_page_break()
    add_heading(document, "How to use one opportunity", 1)
    add_heading(document, "1. Verify the moment", 2)
    add_body(
        document,
        "Open the official source and confirm the organization name, date, location, and the exact event shown. If the record cannot be confirmed, stop and mark it for later research.",
    )
    add_heading(document, "2. Qualify the fit", 2)
    add_body(
        document,
        "Visit the organization's own website. Understand what it does, whether the public event appears relevant, and whether a business-published channel exists. Treat the suggested product fit as a question to investigate, not an answer.",
    )
    add_heading(document, "3. Clear outreach", 2)
    add_body(
        document,
        "Before contact, complete the required suppression, licensing, state-rule, purpose, and channel checks in the protected broker workflow. The public workspace intentionally keeps call-ready at zero.",
    )

    add_heading(document, "Read the record correctly", 2)
    record_table = document.add_table(rows=1, cols=2)
    set_table_geometry(record_table, [2700, 6660])
    mark_header_row(record_table.rows[0])
    for cell in record_table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    set_table_text(record_table.cell(0, 0), "Label", bold=True, color=WHITE)
    set_table_text(record_table.cell(0, 1), "Plain-English meaning", bold=True, color=WHITE)
    rows = [
        ("Source-verified business moment", "The fact returned by the official source."),
        ("Why now", "Why the event may deserve timely research."),
        ("Likely fit", "A business question to investigate, not advice or eligibility."),
        ("Evidence", "The source link, record ID, and receipt attached to the item."),
        ("Limitations", "What the public record does not prove."),
        ("Call ready", "Whether outreach checks have passed. Public view stays at zero."),
    ]
    for index, (label, value) in enumerate(rows, start=1):
        cells = record_table.add_row().cells
        if index % 2 == 0:
            set_cell_shading(cells[0], "F5F8FB")
            set_cell_shading(cells[1], "F5F8FB")
        set_table_text(cells[0], label, bold=True, color=NAVY)
        set_table_text(cells[1], value)
    set_table_geometry(record_table, [2700, 6660])

    document.add_page_break()
    add_heading(document, "Where the live records come from", 1)
    add_body(
        document,
        "The application checks official public sources independently. Each lane reports LIVE, UNAVAILABLE, or NOT APPLICABLE for the selected market.",
    )
    source_table = document.add_table(rows=1, cols=3)
    set_table_geometry(source_table, [3150, 1650, 4560])
    mark_header_row(source_table.rows[0])
    for cell in source_table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    for cell, label in zip(source_table.rows[0].cells, ("Official source", "State at demo time", "What it can show")):
        set_table_text(cell, label, bold=True, color=WHITE, size=10)
    source_rows = [
        ("U.S. Department of Labor Form 5500", "READ RUNTIME", "Benefit-plan filings and reported plan or policy timing fields."),
        ("FMCSA Company Census", "READ RUNTIME", "Organization-level carrier registration activity."),
        ("USAspending", "READ RUNTIME", "Federal contract activity in the current query window."),
        ("EPA ECHO", "READ RUNTIME", "Organization and facility compliance-monitoring activity."),
        ("FCC ULS", "READ RUNTIME", "Organization wireless-license activity when the source answers."),
        ("Chicago business licenses", "READ RUNTIME", "New active business-license records when applicable."),
        ("SAM.gov entity updates", "READ RUNTIME", "Active entity updates when the source and credential are available."),
    ]
    for index, (source, state, meaning) in enumerate(source_rows, start=1):
        cells = source_table.add_row().cells
        if index % 2 == 0:
            for cell in cells:
                set_cell_shading(cell, "F5F8FB")
        set_table_text(cells[0], source, bold=True, color=NAVY, size=10)
        set_table_text(cells[1], state, bold=True, color=TEAL if state == "LIVE" else GOLD, size=10)
        set_table_text(cells[2], meaning, size=10)
    set_table_geometry(source_table, [3150, 1650, 4560])

    add_heading(document, "What is deliberately excluded", 2)
    for text in (
        "Personal social-media profiles, private contacts, and hidden contact enrichment.",
        "Health, age, race, religion, family status, wealth, or insurability inferences.",
        "Named filing signers, administrators, donors, or executives used as personal leads.",
        "Automatic contact permission, automatic quoting, or an underwriting conclusion.",
    ):
        add_list_item(document, text)

    document.add_page_break()
    add_heading(document, "A practical daily routine", 1)
    add_heading(document, "Morning: choose the research queue", 2)
    add_list_item(document, "Open the Broker view and check the release stamp.")
    add_list_item(document, "Choose your state or region; start with one state for depth.")
    add_list_item(document, "Read the Broker brief and pick five organizations to verify.")
    add_list_item(document, "Open every official source before writing a note.")

    add_heading(document, "Research block: turn records into qualified questions", 2)
    add_list_item(document, "Confirm the organization and event on the official record.")
    add_list_item(document, "Check the organization's own website for current operations and a business-published channel.")
    add_list_item(document, "Write down what is known, what is inferred, and what remains unknown.")
    add_list_item(document, "Do not copy sensitive person-level fields into your notes.")

    add_heading(document, "Before any outreach", 2)
    checklist = [
        "The organization and event were verified.",
        "The channel was published by the business itself.",
        "Suppression and Do-Not-Contact checks passed.",
        "Licensing and state rules were checked.",
        "The purpose and product discussion are permitted.",
        "The record does not rely on a sample or unavailable source.",
    ]
    for item in checklist:
        add_list_item(document, item, lead="Check: ")

    add_callout(
        document,
        "If any check is missing:",
        "Do more research or stop. The application is designed to make uncertainty visible, not to push every record into a call list.",
        fill=PALE_GOLD,
    )

    document.add_page_break()
    add_heading(document, "Five-minute meeting walkthrough", 1)
    add_body(
        document,
        "Use this sequence when showing David Leads to a broker, manager, or investor.",
        italic=True,
        color=MUTED,
    )
    add_list_item(document, "Open the app and say: 'This is live organization research from official public records, not a contact list.'", marker="1.")
    add_list_item(document, "Show the 27-state territory selector and focus New York.", marker="2.")
    add_list_item(document, "Point to Live organizations, Source coverage, and Needs research. Explain that unavailable sources stay visible.", marker="3.")
    add_list_item(document, "Open one organization. Read the event, why now, official source, and limitations.", marker="4.")
    add_list_item(document, "Show the receipt or proof link. Explain that the public record can be checked before use.", marker="5.")
    add_list_item(document, "End with the contact gate: research is public; outreach still requires protected human clearance.", marker="6.")

    add_heading(document, "A simple opening line", 2)
    add_callout(
        document,
        "Say:",
        "David Leads helps you decide which organizations to research first, why the timing may matter, and exactly which official record to verify before a compliant conversation.",
        fill=PALE_BLUE,
    )

    add_heading(document, "What not to claim", 2)
    for text in (
        "Do not say every record is a lead ready to call.",
        "Do not say a filing proves need, dissatisfaction, renewal, eligibility, or buying intent.",
        "Do not say a product fit is a recommendation or quote.",
        "Do not say the system uses private profiles or secret contact data.",
        "Do not call a source LIVE if the coverage panel says UNAVAILABLE.",
    ):
        add_list_item(document, text)

    document.add_page_break()
    add_heading(document, "If something looks wrong", 1)
    troubleshooting = [
        ("The page is still opening", "Wait for the opening screen to disappear, then refresh once."),
        ("A source says UNAVAILABLE", "Use another live source or return later. Do not treat the outage as a zero or a live result."),
        ("A state has few records", "Select that one state to run the deeper state-specific pull."),
        ("The count changed", "That is expected. Counts follow current public-source responses and the chosen territory."),
        ("A record looks surprising", "Open the official source and confirm it. If it cannot be confirmed, do not use it."),
        ("You need saved notes or outreach actions", "Those functions belong in the protected broker workflow and are not exposed in the public research view."),
    ]
    trouble_table = document.add_table(rows=1, cols=2)
    set_table_geometry(trouble_table, [3150, 6210])
    mark_header_row(trouble_table.rows[0])
    for cell in trouble_table.rows[0].cells:
        set_cell_shading(cell, NAVY)
    set_table_text(trouble_table.cell(0, 0), "What you see", bold=True, color=WHITE)
    set_table_text(trouble_table.cell(0, 1), "What to do", bold=True, color=WHITE)
    for index, (issue, action) in enumerate(troubleshooting, start=1):
        cells = trouble_table.add_row().cells
        if index % 2 == 0:
            set_cell_shading(cells[0], "F5F8FB")
            set_cell_shading(cells[1], "F5F8FB")
        set_table_text(cells[0], issue, bold=True, color=NAVY)
        set_table_text(cells[1], action)
    set_table_geometry(trouble_table, [3150, 6210])

    add_heading(document, "The one rule to remember", 1)
    add_callout(
        document,
        "Use David Leads to decide what to research first.",
        "Use the official record to verify why. Use the protected clearance process to decide whether contact is permitted.",
    )
    add_body(
        document,
        "Public organization data only. No passwords or access keys are printed in this guide.",
        italic=True,
        color=MUTED,
    )

    return document


if __name__ == "__main__":
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build()
    document.save(OUTPUT)
    print(f"written {OUTPUT.resolve()} ({OUTPUT.stat().st_size} bytes)")
