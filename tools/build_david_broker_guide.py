from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "David_Leads_Broker_Guide.docx"
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
TEAL = RGBColor(20, 137, 132)
INK = RGBColor(28, 39, 55)
MUTED = RGBColor(92, 105, 122)
PALE_BLUE = "E8EEF5"
PALE_TEAL = "E7F5F3"
PALE_GOLD = "FFF7E5"


def _ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def _set_width(parent, tag: str, width_dxa: int) -> None:
    width = _ensure_child(parent, tag)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(int(width_dxa)))


def apply_table_geometry(
    table,
    column_widths_dxa,
    *,
    table_width_dxa: int,
    indent_dxa: int,
    cell_margins_dxa: dict[str, int],
) -> None:
    """Apply deterministic table width, grid, cell width, and margins."""

    widths = [int(width) for width in column_widths_dxa]
    if not widths or any(width <= 0 for width in widths):
        raise ValueError("table column widths must be positive")
    if sum(widths) != int(table_width_dxa):
        raise ValueError("table column widths must sum to table_width_dxa")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    _set_width(table_properties, "w:tblW", table_width_dxa)
    table_indent = _ensure_child(table_properties, "w:tblInd")
    table_indent.set(qn("w:type"), "dxa")
    table_indent.set(qn("w:w"), str(int(indent_dxa)))
    layout = _ensure_child(table_properties, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for column_index, width in enumerate(widths):
        table.columns[column_index].width = Twips(width)
    for row in table.rows:
        if len(row.cells) != len(widths):
            raise ValueError("table rows must be unmerged before geometry is applied")
        for column_index, cell in enumerate(row.cells):
            width = widths[column_index]
            cell.width = Twips(width)
            cell_properties = cell._tc.get_or_add_tcPr()
            _set_width(cell_properties, "w:tcW", width)
            margins = _ensure_child(cell_properties, "w:tcMar")
            for side in ("top", "bottom", "start", "end"):
                margin = _ensure_child(margins, f"w:{side}")
                margin.set(qn("w:w"), str(int(cell_margins_dxa[side])))
                margin.set(qn("w:type"), "dxa")


def set_font(run, *, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def shade_paragraph(paragraph, fill: str, border: str | None = None):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_text(doc, text, *, bold=False, color=INK, size=11, after=6, italic=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    set_font(
        paragraph.add_run(text),
        size=size,
        bold=bold,
        color=color,
        italic=italic,
    )
    return paragraph


def add_callout(doc, label: str, text: str, *, fill=PALE_TEAL, border="148984"):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    shade_paragraph(paragraph, fill, border)
    set_font(paragraph.add_run(f"{label}: "), bold=True, color=NAVY)
    set_font(paragraph.add_run(text), color=INK)
    return paragraph


def add_bullet(doc, text: str, *, bold_lead: str | None = None):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_lead and text.startswith(bold_lead):
        set_font(paragraph.add_run(bold_lead), bold=True, color=NAVY)
        set_font(paragraph.add_run(text[len(bold_lead):]), color=INK)
    else:
        set_font(paragraph.add_run(text), color=INK)
    return paragraph


def add_number(doc, number: int, title: str, detail: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.25
    set_font(paragraph.add_run(f"{number}. "), bold=True, color=NAVY)
    set_font(paragraph.add_run(f"{title}. "), bold=True, color=NAVY)
    set_font(paragraph.add_run(detail), color=INK)
    return paragraph


def set_cell_fill(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def add_metadata_grid(doc):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    values = [
        ("OPEN", "No login for the public research view"),
        ("COVERAGE", "27 Eastern U.S. markets"),
        ("PRIMARY USE", "Choose what to research first"),
        ("CONTACT", "Always blocked until protected clearance"),
    ]
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for cell, text in zip(header.cells, ("AT A GLANCE", "WHAT IT MEANS")):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(cell, "0B2545")
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        set_font(paragraph.add_run(text), size=8.5, bold=True, color=RGBColor(255, 255, 255))
    for index, (label, value) in enumerate(values, start=1):
        left, right = table.rows[index].cells
        left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_fill(left, PALE_BLUE)
        set_cell_fill(right, "FFFFFF")
        paragraph = left.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.1
        set_font(paragraph.add_run(label), size=8.5, bold=True, color=TEAL)
        detail = right.paragraphs[0]
        detail.paragraph_format.space_before = Pt(0)
        detail.paragraph_format.space_after = Pt(0)
        detail.paragraph_format.line_spacing = 1.15
        set_font(detail.add_run(value), size=10.5, bold=True, color=NAVY)
    apply_table_geometry(
        table,
        [2200, 7160],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 130, "bottom": 130, "start": 120, "end": 120},
    )
    return table


def add_page_number(paragraph):
    set_font(paragraph.add_run("David Leads | Broker Guide   •   Page "), size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = OxmlElement("w:r")
    run.append(begin)
    run.append(instruction)
    run.append(separate)
    run.append(value)
    run.append(end)
    paragraph._p.append(run)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, RGBColor(31, 77, 120), 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    set_font(header_paragraph.add_run("DAVID LEADS"), size=8.5, bold=True, color=TEAL)
    set_font(header_paragraph.add_run("   |   Evidence-Backed Broker Research"), size=8.5, color=MUTED)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_paragraph.paragraph_format.space_before = Pt(0)
    footer_paragraph.paragraph_format.space_after = Pt(0)
    add_page_number(footer_paragraph)


def heading(doc, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    keep_with_next(paragraph)
    return paragraph


def build():
    doc = Document()
    configure_document(doc)
    doc.core_properties.title = "David Leads - Broker Guide"
    doc.core_properties.subject = "Plain-language operating guide for David Abraham"
    doc.core_properties.author = "SZL Holdings"
    doc.core_properties.keywords = "David Leads, broker guide, public data, life insurance"

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(10)
    kicker.paragraph_format.space_after = Pt(2)
    set_font(kicker.add_run("BROKER OPERATING GUIDE"), size=9, bold=True, color=TEAL)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("David Leads"), size=30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_font(
        subtitle.add_run("From source-verified business signal to the next lawful broker action"),
        size=14,
        color=MUTED,
    )

    add_metadata_grid(doc)
    add_text(
        doc,
        "Open the application",
        bold=True,
        color=TEAL,
        size=9,
        after=2,
    )
    add_text(
        doc,
        "https://szlholdings-david-leads.hf.space/",
        bold=True,
        color=BLUE,
        size=12,
        after=12,
    )
    add_callout(
        doc,
        "The simple rule",
        "Use David Leads to decide what to research first. Use the official record "
        "to verify why. Use the protected clearance process to decide whether contact "
        "is permitted.",
    )

    heading(doc, "What this application does", 1)
    add_text(
        doc,
        "David Leads organizes official organization-level records into a practical "
        "research queue. It shows the company, the observed change, the timing, the "
        "likely protection-planning fit, and the evidence behind the item.",
    )
    add_callout(
        doc,
        "Important",
        "It is not a purchased contact list. It never turns a public record into "
        "permission to call, email, market, or make an underwriting decision.",
        fill=PALE_GOLD,
        border="B78921",
    )

    doc.add_page_break()
    heading(doc, "Start your day in five minutes", 1)
    add_number(doc, 1, "Read the Broker brief", "It summarizes the strongest timing pattern in the current official-source pull.")
    add_number(doc, 2, "Choose a territory", "Select All East, a region, or one state. One state runs a deeper market-specific search.")
    add_number(doc, 3, "Choose a deal moment", "Use Life-plan timing, Growth & awards, Operational change, or Needs research.")
    add_number(doc, 4, "Open the strongest account", "Review the timing, likely fit, evidence, limitations, and three-step broker action.")
    add_number(doc, 5, "Verify before doing anything else", "Open the cited official record and confirm that the organization and event are current.")

    heading(doc, "The four deal-moment lanes", 1)
    add_bullet(
        doc,
        "Life-plan timing: Employers whose official Department of Labor filing "
        "reports a group-life benefit and a plan or policy anniversary.",
        bold_lead="Life-plan timing:",
    )
    add_bullet(
        doc,
        "Growth & awards: Current federal contract activity that may justify a "
        "business-protection, continuity, or workforce review.",
        bold_lead="Growth & awards:",
    )
    add_bullet(
        doc,
        "Operational change: Carrier registrations, facility monitoring activity, "
        "and other official organization events.",
        bold_lead="Operational change:",
    )
    add_bullet(
        doc,
        "Needs research: Every account that still requires source verification, "
        "product-fit review, and protected contact clearance.",
        bold_lead="Needs research:",
    )
    add_callout(
        doc,
        "Timing label",
        "A Form 5500 anniversary is a research hypothesis. It is not proof of a "
        "renewal, dissatisfaction, buying intent, eligibility, or insurability.",
    )

    doc.add_page_break()
    heading(doc, "How to read an account", 1)
    heading(doc, "Source-verified business moment", 2)
    add_text(doc, "The organization-level fact returned by the official source. Read this before the product idea.")
    heading(doc, "Timing", 2)
    add_text(doc, "The observed date or the next anniversary calculated from a previously reported plan or policy period.")
    heading(doc, "Likely fit", 2)
    add_text(
        doc,
        "A reason to investigate protection planning. It is not a recommendation, "
        "quote, coverage determination, or statement that David is appointed or "
        "licensed for every product shown by the source.",
    )
    heading(doc, "Evidence", 2)
    add_text(
        doc,
        "A direct filing or official record shows the evidence class. A session-verifiable "
        "source reference binds the normalized source record to its receipt. Identity "
        "resolution, proof grade, clock, counter-evidence, and durable history remain "
        "outside that receipt and must be evaluated separately.",
    )
    heading(doc, "Contact state", 2)
    add_text(
        doc,
        "Cleared to contact stays at zero in the public view. That is intentional. "
        "Public visibility does not create consent.",
    )

    heading(doc, "The three-step broker action", 1)
    add_number(doc, 1, "Verify the moment", "Open the cited source and confirm the organization, reported date, and relevant filing field.")
    add_number(doc, 2, "Qualify the fit", "Use the organization’s own website to understand its operations and find a business-published channel.")
    add_number(doc, 3, "Clear outreach", "In the protected workflow, complete suppression, licensing, state-rule, purpose, and channel checks before any call or email.")

    doc.add_page_break()
    heading(doc, "What David Leads never assumes", 1)
    add_bullet(doc, "A public record is not consent or permission to contact.")
    add_bullet(doc, "Observed activity is not revenue, intent, insurability, or a deal value.")
    add_bullet(doc, "The system does not infer health, age, race, religion, family status, or private wealth.")
    add_bullet(doc, "The system does not scrape social profiles or use named executives, insiders, signers, or donors as personal wealth leads.")
    add_bullet(doc, "The public view never exposes broker notes, saved channels, clearance decisions, outcomes, or exports.")

    heading(doc, "Department of Labor privacy boundary", 1)
    add_text(
        doc,
        "The life-plan lane keeps only the organization, public location, plan-period "
        "dates, participant count, reported carriers, and benefit categories needed "
        "for the research card.",
    )
    add_callout(
        doc,
        "Excluded",
        "EINs, phone numbers, named administrators and signers, preparers, broker "
        "identities, commissions, and person-level addresses are not retained.",
        fill=PALE_GOLD,
        border="B78921",
    )

    heading(doc, "When a source is unavailable", 1)
    add_number(doc, 1, "Open Market coverage", "Each source reports LIVE, UNAVAILABLE, or NOT APPLICABLE.")
    add_number(doc, 2, "Read the reason", "A missing credential, source outage, or territory mismatch remains visible.")
    add_number(doc, 3, "Do not work around it", "The application never substitutes a fake lead for an unavailable source.")

    heading(doc, "Best practice before every outreach decision", 1)
    add_bullet(doc, "Confirm the official record is current.")
    add_bullet(doc, "Confirm the product and jurisdiction are within David’s approved scope.")
    add_bullet(doc, "Use only a business-published channel from the organization’s own website.")
    add_bullet(doc, "Complete all suppression, licensing, state-rule, and purpose checks.")
    add_bullet(doc, "Record time-limited clearance in the protected workflow.")

    add_callout(
        doc,
        "Remember",
        "The software helps David focus his research. David’s licensed judgment and "
        "the protected clearance process decide whether any next action is appropriate.",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
